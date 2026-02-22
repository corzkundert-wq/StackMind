import os
import uuid
import logging
from sqlalchemy.orm import Session as DBSession
from src.backend.models import File, Chunk, FileStatus
from src.backend.services.llm_service import get_embeddings

logger = logging.getLogger("stackmind")
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def extract_text(filepath: str, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    try:
        if ext in ("txt", "md", "csv", "json", "py", "js", "html", "css", "xml", "log", "yaml", "yml", "toml", "ini", "cfg", "rst", "tsv"):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        elif ext == "docx":
            try:
                from docx import Document
                doc = Document(filepath)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        tables_text.append(" | ".join(row_data))
                return "\n\n".join(paragraphs) + ("\n\nTables:\n" + "\n".join(tables_text) if tables_text else "")
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                with open(filepath, "rb") as f:
                    return f.read().decode("utf-8", errors="ignore")

        elif ext == "doc":
            try:
                import subprocess
                result = subprocess.run(["strings", filepath], capture_output=True, text=True, timeout=30)
                text = result.stdout
                lines = [l.strip() for l in text.split("\n") if len(l.strip()) > 3]
                return "\n".join(lines)
            except Exception as e:
                logger.error(f"DOC extraction failed: {e}")
                with open(filepath, "rb") as f:
                    return f.read().decode("latin-1", errors="ignore")[:50000]

        elif ext in ("xlsx", "xls"):
            try:
                from openpyxl import load_workbook
                wb = load_workbook(filepath, read_only=True, data_only=True)
                all_text = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    all_text.append(f"--- Sheet: {sheet_name} ---")
                    for row in ws.iter_rows(values_only=True):
                        row_vals = [str(c) if c is not None else "" for c in row]
                        if any(v.strip() for v in row_vals):
                            all_text.append(" | ".join(row_vals))
                wb.close()
                return "\n".join(all_text)
            except Exception as e:
                logger.error(f"Excel extraction failed: {e}")
                return f"Excel file: {filename} (extraction failed: {str(e)})"

        elif ext == "pdf":
            try:
                with open(filepath, "rb") as f:
                    content = f.read()
                text_parts = []
                raw = content.decode("latin-1", errors="ignore")
                import re
                streams = re.findall(r'BT\s(.*?)ET', raw, re.DOTALL)
                for stream in streams:
                    text_matches = re.findall(r'\((.*?)\)', stream)
                    text_parts.extend(text_matches)
                extracted = " ".join(text_parts).strip()
                if len(extracted) > 50:
                    return extracted[:100000]
                return raw[:50000]
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                with open(filepath, "rb") as f:
                    return f.read().decode("latin-1", errors="ignore")[:50000]

        elif ext in ("wav", "mp3", "m4a", "ogg", "webm", "flac", "aac"):
            return f"[Audio file: {filename} — use transcription to extract text]"

        else:
            try:
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception:
                with open(filepath, "rb") as f:
                    return f.read().decode("latin-1", errors="ignore")[:50000]

    except Exception as e:
        logger.error(f"Text extraction failed for {filename}: {e}")
        raise


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list:
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap
        if start >= len(text):
            break
    return chunks if chunks else [text[:chunk_size]]


def process_file(db: DBSession, file_record: File, filepath: str):
    try:
        file_record.status = FileStatus.uploaded
        db.commit()

        raw_text = extract_text(filepath, file_record.filename)
        file_record.raw_text = raw_text
        file_record.status = FileStatus.extracted
        db.commit()

        chunks = chunk_text(raw_text)
        for i, chunk_text_content in enumerate(chunks):
            chunk = Chunk(
                file_id=file_record.id,
                chunk_index=i,
                text=chunk_text_content,
                chunk_metadata={"char_count": len(chunk_text_content)},
            )
            db.add(chunk)
        file_record.status = FileStatus.chunked
        db.commit()

        db_chunks = db.query(Chunk).filter(Chunk.file_id == file_record.id).order_by(Chunk.chunk_index).all()
        chunk_texts = [c.text for c in db_chunks]
        embeddings = get_embeddings(chunk_texts)
        for chunk_obj, emb in zip(db_chunks, embeddings):
            chunk_obj.chunk_metadata = {**(chunk_obj.chunk_metadata or {}), "embedding": emb}
        file_record.status = FileStatus.embedded
        db.commit()

    except Exception as e:
        logger.error(f"File processing failed: {e}")
        file_record.status = FileStatus.failed
        file_record.error = str(e)
        db.commit()
        raise


def process_text_paste(db: DBSession, library_id: str, title: str, text: str, tags: list = None) -> File:
    safe_name = f"{uuid.uuid4().hex}_{title.replace(' ', '_')}.txt"
    filepath = os.path.join(UPLOAD_DIR, safe_name)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(text)

    file_record = File(
        library_id=library_id,
        filename=safe_name,
        display_name=title or "Pasted Text",
        tags=tags or ["pasted"],
        status=FileStatus.uploaded,
    )
    db.add(file_record)
    db.commit()
    db.refresh(file_record)

    process_file(db, file_record, filepath)
    db.refresh(file_record)
    return file_record


def save_upload(file_bytes: bytes, filename: str) -> str:
    safe_name = f"{uuid.uuid4().hex}_{filename}"
    filepath = os.path.join(UPLOAD_DIR, safe_name)
    with open(filepath, "wb") as f:
        f.write(file_bytes)
    return filepath
